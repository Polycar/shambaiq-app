import os
import win32com.client
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image, ImageDraw, ImageFont

def generate_logo():
    print("Generating ShambaIQ brand logo/header PNG...")
    # Create a clean high-res canvas (1200 x 300) for high print quality
    width, height = 1200, 300
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0)) # Transparent background
    draw = ImageDraw.Draw(image)

    # Brand Colors
    c_dark_green = (26, 58, 26, 255)    # #1a3a1a
    c_brand_green = (22, 163, 74, 255)  # #16a34a
    c_gold = (200, 134, 10, 255)        # #C8860A
    c_brown = (139, 105, 20, 255)       # #8B6914

    # Draw a stylized icon (Concept 2 - Pin & Sprout inspired) on the left (x: 50, y: 50)
    # Map pin shape
    # We will draw a custom path or overlapping geometric shapes
    # Pin body
    pin_x, pin_y = 100, 150
    draw.ellipse([pin_x - 45, pin_y - 85, pin_x + 45, pin_y + 5], fill=c_brand_green)
    # Triangle bottom of pin
    draw.polygon([pin_x - 38, pin_y - 20, pin_x, pin_y + 40, pin_x + 38, pin_y - 20], fill=c_brand_green)
    # Outer white/cream border ring inside
    draw.ellipse([pin_x - 30, pin_y - 70, pin_x + 30, pin_y - 10], fill=(255, 255, 255, 255))
    # Plant sprout inside pin
    # Stem
    draw.line([pin_x, pin_y - 15, pin_x, pin_y - 45], fill=c_brand_green, width=6)
    # Left Leaf
    draw.ellipse([pin_x - 22, pin_y - 50, pin_x, pin_y - 35], fill=c_brand_green)
    # Right Leaf
    draw.ellipse([pin_x, pin_y - 55, pin_x + 22, pin_y - 40], fill=c_gold)
    
    # Soil dots at pin base
    draw.ellipse([pin_x - 55, pin_y + 45, pin_x - 50, pin_y + 50], fill=c_brown)
    draw.ellipse([pin_x - 15, pin_y + 50, pin_x - 10, pin_y + 55], fill=c_brown)
    draw.ellipse([pin_x + 25, pin_y + 48, pin_x + 30, pin_y + 53], fill=c_brown)

    # Load Fonts
    def get_font(name, size):
        try:
            paths = [
                f"C:\\Windows\\Fonts\\{name}.ttf",
                f"C:\\Windows\\Fonts\\{name.lower()}.ttf",
                name
            ]
            for p in paths:
                if os.path.exists(p):
                    return ImageFont.truetype(p, size)
            # Default fallback search for common fonts
            return ImageFont.truetype("georgiab.ttf", size)
        except IOError:
            try:
                return ImageFont.truetype("arial.ttf", size)
            except IOError:
                return ImageFont.load_default()

    font_title = get_font("georgiab", 84) # Georgia Bold
    font_subtitle = get_font("arial", 32) # Arial

    # Write text "ShambaIQ"
    draw.text((180, 50), "Shamba", fill=c_dark_green, font=font_title)
    draw.text((510, 50), "IQ", fill=c_gold, font=font_title)

    # Tagline text
    draw.text((185, 160), "Precision Agriculture & Soil Intelligence", fill=c_brown, font=font_subtitle)

    # Ensure public folder exists
    os.makedirs("public", exist_ok=True)
    logo_path = os.path.abspath("public/shambaiq_logo_header.png")
    image.save(logo_path, "PNG")
    print(f"Logo generated at {logo_path}")
    return logo_path

def add_header_footer(doc):
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = "ShambaIQ B2B Operations  |  Email: info@shambaiq.com  |  Phone: +254 748 042 633"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(128, 128, 128)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_paragraph_callout(paragraph, bg_color_hex, border_color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="15" w:color="{border_color_hex}"/></w:pBdr>')
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color_hex}"/>')
    pPr.append(pbdr)
    pPr.append(shd)

def create_partnership_doc(logo_path):
    print("Building ShambaIQ B2B Partnership & Licensing Suite DOCX...")
    doc = Document()
    add_header_footer(doc)

    # 1. Logo
    doc.add_picture(logo_path, width=Inches(5.0))
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # 2. Title & Subtitle
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("B2B Partnership & Licensing Suite")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 58, 26) # Dark Green

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Precision Agriculture API & Dynamic Soil Intelligence for Kenya")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(200, 134, 10) # Gold

    # 3. Metadata box
    p_meta = doc.add_paragraph()
    set_paragraph_callout(p_meta, "F5F0E1", "8B6914") # Cream bg, brown border
    p_meta.paragraph_format.space_after = Pt(24)
    
    meta_text = (
        "Document Version: 1.0 (May 2026)\n"
        "Prepared by: ShambaIQ B2B Licensing & Operations\n"
        "Official Email: info@shambaiq.com  |  Official WhatsApp: +254 748 042 633\n"
        "Principal Office: Nairobi, Kenya"
    )
    run_meta = p_meta.add_run(meta_text)
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(26, 58, 26)

    # 4. Heading 1: Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run_h1 = h1.add_run("1. Executive Summary")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = RGBColor(26, 58, 26)

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(12)
    run_exec = p_exec.add_run(
        "ShambaIQ is a state-of-the-art agronomic decision-support platform designed specifically for the East African agricultural landscape. "
        "By linking high-resolution satellite data (iSDAsoil) with local crop suitability models and real-time input pricing, "
        "ShambaIQ delivers hyper-local, chemically exact soil diagnostics and fertilizer recommendations.\n\n"
        "We offer non-exclusive B2B software licenses and secure API integrations to cooperatives, microfinance organizations, "
        "non-governmental organizations (NGOs), and agricultural input brands. Partnering with ShambaIQ allows your organization "
        "to offer world-class, data-driven agricultural advisory directly to your users—boosting crop yields, reducing agricultural "
        "loan risks, and driving regional input sales."
    )
    run_exec.font.name = 'Arial'
    run_exec.font.size = Pt(11)
    run_exec.font.color.rgb = RGBColor(51, 51, 51)

    # 5. Heading 1: The ShambaIQ Advantage
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    run_h2 = h2.add_run("2. The ShambaIQ Advantage")
    run_h2.font.name = 'Georgia'
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = RGBColor(26, 58, 26)

    advantages = [
        ("Sub-Second Dynamic Soil Profiles: ", "Access spatial soil diagnostics (pH, Nitrogen, Phosphorus, Potassium) for all 47 counties and deep ward-level zones across Kenya."),
        ("Instant, Low-Data Delivery: ", "Pre-compiled server-side technology ensures pages load in milliseconds, even on unstable rural 3G/4G connections."),
        ("Zero-Disclosure Security (Black-Box API): ", "Our proprietary recommendation engine resides entirely on our secure servers. Your system queries our endpoints via secure API tokens and receives clean JSON outputs. Your developers never need to rebuild or host complex agronomic databases."),
        ("Search Engine Dominance: ", "ShambaIQ boasts an active SEO footprint of over 2,700 high-authority, hyper-local dynamic pages, capturing organic farmer traffic at the precise moment they seek inputs.")
    ]

    for title, desc in advantages:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(6)
        r_title = p_bullet.add_run(title)
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(26, 58, 26)
        r_desc = p_bullet.add_run(desc)
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(11)
        r_desc.font.color.rgb = RGBColor(51, 51, 51)

    # Page Break before pricing
    doc.add_page_break()

    # 6. Heading 1: Licensing Tiers & Commercial Packages
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)
    run_h3 = h3.add_run("3. Licensing Tiers & Commercial Packages")
    run_h3.font.name = 'Georgia'
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = RGBColor(26, 58, 26)

    doc.add_paragraph("We offer three standard commercial packages designed to fit varying organization sizes and technical capabilities:").runs[0].font.name = 'Arial'

    # Create Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Package / Licensing Tier'
    hdr_cells[1].text = 'Best For & Technical Setup'
    hdr_cells[2].text = 'Commercial Pricing (KES)'

    # Style Header row
    for cell in hdr_cells:
        set_cell_shading(cell, "16A34A")
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10.5)

    # Row 1
    row_cells = table.rows[1].cells
    row_cells[0].text = "Tier 1: Precision Advisory API (B2B SaaS)"
    row_cells[1].text = "Mobile AgTech platforms, SMS/WhatsApp chatbot operators, and digital microfinance providers. Integration via standard REST API."
    row_cells[2].text = "• Startup: KES 30,000 / month (Up to 15k queries)\n• Enterprise: KES 75,000 / month (Up to 100k queries)\n• Volume: KES 0.65 per query thereafter"

    # Row 2
    row_cells = table.rows[2].cells
    row_cells[0].text = "Tier 2: Co-Branded Portal License"
    row_cells[1].text = "Agricultural cooperatives, county governments, and regional developmental NGOs. Co-branded portal hosted on coop.shambaiq.com."
    row_cells[2].text = "• Annual Base License: KES 250,000 / year (hosting, SSL, maintenance)\n• Setup & Customization: KES 50,000 (one-time)"

    # Row 3
    row_cells = table.rows[3].cells
    row_cells[0].text = "Tier 3: Verified Agrovet & Input Placement"
    row_cells[1].text = "Seed producers, fertilizer manufacturers, and regional agrovets. Banner ads and dealer links on high-converting soil report pages."
    row_cells[2].text = "• County-Level Placement: KES 5,000 / month per county guide\n• National Placement: KES 120,000 / month (all 47 counties)"

    # Style table text
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue
        # Alternating row colors
        bg_color = "F5F0E1" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg_color)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 51, 51)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(18)

    # 7. Heading 1: API Request & Response Specification
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(8)
    run_h4 = h4.add_run("4. API Request & Response Specification")
    run_h4.font.name = 'Georgia'
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = RGBColor(26, 58, 26)

    doc.add_paragraph("Integrate ShambaIQ advisory cards into your application with standard JSON endpoints:").runs[0].font.name = 'Arial'

    # Sample Request Code Box
    doc.add_paragraph("Sample HTTP GET Request:").runs[0].font.bold = True
    p_req = doc.add_paragraph()
    set_paragraph_callout(p_req, "F5F5F5", "16A34A")
    r_req = p_req.add_run("GET https://api.shambaiq.com/v1/recommend?county=Kiambu&crop=Cabbage&key=YOUR_API_TOKEN")
    r_req.font.name = 'Consolas'
    r_req.font.size = Pt(9.5)
    r_req.font.color.rgb = RGBColor(26, 58, 26)

    # Sample Response Code Box
    doc.add_paragraph("Sample JSON Response Payload:").runs[0].font.bold = True
    p_resp = doc.add_paragraph()
    set_paragraph_callout(p_resp, "F5F5F5", "16A34A")
    
    json_payload = (
        "{\n"
        "  \"status\": \"success\",\n"
        "  \"query\": {\n"
        "    \"county\": \"Kiambu\",\n"
        "    \"crop\": \"Cabbage\"\n"
        "  },\n"
        "  \"soil_diagnostics\": {\n"
        "    \"average_pH\": 5.28,\n"
        "    \"classification\": \"Moderately Acidic\",\n"
        "    \"nitrogen_g_kg\": 0.98,\n"
        "    \"extractable_p_mg_kg\": 18.2\n"
        "  },\n"
        "  \"prescriptions\": {\n"
        "    \"basal_dressing\": \"Apply 75kg/acre of buffered NPK planting fertilizer. Avoid raw DAP to prevent phosphorus locking in acidic soils.\",\n"
        "    \"soil_adjustment\": \"Incorporate 1.5 tons/acre of agricultural lime during land preparation to buffer soil pH above 6.0.\",\n"
        "    \"top_dressing\": \"Apply 50kg/acre of Calcium Ammonium Nitrate (CAN) at week 3 and week 6 post-transplant.\"\n"
        "  },\n"
        "  \"links\": {\n"
        "    \"advisory_hub\": \"https://www.shambaiq.com/soil/kiambu/cabbage\"\n"
        "  }\n"
        "}"
    )
    r_resp = p_resp.add_run(json_payload)
    r_resp.font.name = 'Consolas'
    r_resp.font.size = Pt(9.0)
    r_resp.font.color.rgb = RGBColor(51, 51, 51)

    # 8. Heading 1: Get Started & IP Security
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(8)
    run_h5 = h5.add_run("5. Get Started")
    run_h5.font.name = 'Georgia'
    run_h5.font.size = Pt(16)
    run_h5.font.bold = True
    run_h5.font.color.rgb = RGBColor(26, 58, 26)

    p_start = doc.add_paragraph()
    p_start.paragraph_format.space_after = Pt(12)
    run_start = p_start.add_run(
        "Protecting our collaborative intellectual property is a core priority. To request a sandbox API key or schedule a white-label portal demo, "
        "please review our standard Mutual NDA and contact our licensing team using the details below. "
        "We are excited to build the future of precision agriculture with your team."
    )
    run_start.font.name = 'Arial'
    run_start.font.size = Pt(11)
    run_start.font.color.rgb = RGBColor(51, 51, 51)

    # Callout for contact info
    p_contact = doc.add_paragraph()
    set_paragraph_callout(p_contact, "E8F5E9", "16A34A") # Soft green bg, green border
    p_contact.paragraph_format.space_after = Pt(12)
    
    contact_text = (
        "SHAMBAIQ LICENSING & B2B OPERATIONS\n"
        "• Primary Email: info@shambaiq.com\n"
        "• WhatsApp Business Support: +254 748 042 633\n"
        "• Operations Center: Nairobi, Kenya"
    )
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = 'Arial'
    run_contact.font.size = Pt(11)
    run_contact.font.bold = True
    run_contact.font.color.rgb = RGBColor(26, 58, 26)

    # Save DOCX
    docx_path = os.path.abspath("shambaiq_partnership_proposal.docx")
    doc.save(docx_path)
    print(f"Partnership Suite DOCX saved to {docx_path}")
    return docx_path

def create_nda_doc(logo_path):
    print("Building ShambaIQ Mutual NDA DOCX...")
    doc = Document()
    add_header_footer(doc)

    # 1. Logo
    doc.add_picture(logo_path, width=Inches(5.0))
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # 2. Title & Subtitle
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("MUTUAL NON-DISCLOSURE AGREEMENT")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 58, 26) # Dark Green

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Proprietary & Intellectual Property Protection Agreement")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(200, 134, 10) # Gold

    # 3. Preamble
    p_preamble = doc.add_paragraph()
    p_preamble.paragraph_format.space_after = Pt(12)
    r_pre = p_preamble.add_run(
        "THIS AGREEMENT is entered into as of this _____ day of May, 2026 (the \"Effective Date\"), by and between:\n\n"
        "1. SHAMBAIQ OPERATIONS LIMITED, a private limited company incorporated under the laws of the Republic of Kenya, "
        "with its principal office in Nairobi, Kenya (contact: info@shambaiq.com, +254 748 042 633) (hereinafter \"ShambaIQ\"); and\n"
        "2. THE PARTNER ORGANIZATION (as identified in the signature block below, hereinafter \"Partner\").\n\n"
        "ShambaIQ and the Partner are collectively referred to as the \"Parties\" and individually as a \"Party.\""
    )
    r_pre.font.name = 'Arial'
    r_pre.font.size = Pt(11)
    r_pre.font.color.rgb = RGBColor(51, 51, 51)

    # Section generator function
    sections_data = [
        ("1. Purpose of Agreement", 
         "The Parties wish to engage in exploratory discussions and evaluations regarding a potential business relationship, specifically concerning ShambaIQ's B2B Software Licensing, API Integrations, agricultural decision-support algorithms, and county-level soil intelligence platforms (the \"Purpose\"). In connection with this Purpose, either Party (the \"Disclosing Party\") may disclose proprietary, technical, financial, or commercial information to the other Party (the \"Receiving Party\")."),
        
        ("2. Definition of Confidential Information", 
         "\"Confidential Information\" refers to any information, in whatever form, disclosed by the Disclosing Party to the Receiving Party that is marked as confidential or that should reasonably be understood to be confidential given the nature of the information or the circumstances of disclosure.\n\n"
         "Without limiting the generality of the foregoing, Confidential Information includes:\n"
         "• For ShambaIQ: Agronomic algorithms, database structures, iSDAsoil integrations, crop suitability matrices, API request/response payloads, pricing models, source code, and county soil diagnostics data.\n"
         "• For the Partner: Business operations data, customer lists, marketing strategies, system integrations, and SMS/USSD platform telemetry.\n"
         "• For Both: The existence, status, and terms of any discussions, negotiations, or potential agreements between the Parties."),
        
        ("3. Obligations of the Receiving Party", 
         "The Receiving Party shall:\n"
         "1. Maintain Strict Secrecy: Keep all Confidential Information strictly confidential and apply at least the same degree of care it uses to protect its own confidential information of like importance, but in no event less than a reasonable standard of care.\n"
         "2. Limit Purpose: Use the Confidential Information solely for the Purpose of evaluating and executing the potential business relationship described in Section 1.\n"
         "3. Restrict Access: Disclose the Confidential Information only to those of its directors, officers, employees, or legal advisers who have a clear \"need to know\" for the Purpose, and who are bound by confidentiality obligations no less restrictive than those in this Agreement.\n"
         "4. Prevent Copying: Not copy, reproduce, reverse engineer, decompile, or disassemble any software, databases, or algorithms provided by the Disclosing Party without prior written consent."),

        ("4. Exclusions from Confidentiality", 
         "This Agreement imposes no obligation upon the Receiving Party with respect to information that:\n"
         "1. Is or becomes publicly known through no breach of this Agreement by the Receiving Party;\n"
         "2. Was already in the rightful possession of the Receiving Party prior to disclosure, without confidentiality restrictions;\n"
         "3. Is independently developed by the Receiving Party without reference to or reliance upon the Disclosing Party’s Confidential Information;\n"
         "4. Is rightfully obtained by the Receiving Party from a third party without restrictions on disclosure."),

        ("5. Compelled Disclosure", 
         "If the Receiving Party is required by law, regulation, or a court of competent jurisdiction to disclose any Confidential Information, it shall, to the extent legally permitted, provide the Disclosing Party with prompt written notice so that the Disclosing Party may seek a protective order or other appropriate remedy."),

        ("6. Term and Termination", 
         "This Agreement shall commence on the Effective Date and shall remain in effect for a period of two (2) years, unless terminated earlier by either Party upon thirty (30) days' written notice.\n\n"
         "The confidentiality obligations set forth herein with respect to any Confidential Information disclosed during the term shall survive the termination or expiration of this Agreement for a period of three (3) years from the date of disclosure."),

        ("7. Return or Destruction of Materials", 
         "Upon written request by the Disclosing Party, or upon termination of this Agreement, the Receiving Party shall immediately return to the Disclosing Party or destroy (and certify such destruction in writing) all documents, notes, models, code repositories, or other physical or electronic media containing or reflecting any Confidential Information."),

        ("8. Intellectual Property & Warranties", 
         "1. No License: No license, patent, copyright, trademark, or other intellectual property right is granted, implied, or transferred under this Agreement. All Confidential Information remains the exclusive property of the Disclosing Party.\n"
         "2. As-Is Basis: All Confidential Information is provided on an \"AS IS\" basis. Neither Party makes any warranty, express or implied, as to its accuracy, completeness, or suitability for any particular purpose."),

        ("9. Governing Law and Dispute Resolution", 
         "This Agreement shall be governed by and construed in accordance with the Laws of the Republic of Kenya.\n\n"
         "Any dispute, controversy, or claim arising out of or in connection with this Agreement, including its existence, validity, or termination, shall be resolved amicably through good-faith negotiations. If negotiations fail within thirty (30) days, the dispute shall be referred to and finally resolved by arbitration under the rules of the Nairobi Centre for International Arbitration (NCIA). The place of arbitration shall be Nairobi, and the language shall be English."),

        ("10. Miscellaneous Provisions", 
         "1. Entire Agreement: This Agreement constitutes the entire understanding between the Parties concerning the subject matter and supersedes all prior discussions.\n"
         "2. Amendments: No amendment, modification, or waiver of any provision of this Agreement shall be effective unless in writing and signed by authorized representatives of both Parties.\n"
         "3. Severability: If any provision of this Agreement is held to be invalid or unenforceable, the remaining provisions shall remain in full force and effect.")
    ]

    for title, text in sections_data:
        h_sec = doc.add_paragraph()
        h_sec.paragraph_format.space_before = Pt(14)
        h_sec.paragraph_format.space_after = Pt(6)
        r_t = h_sec.add_run(title)
        r_t.font.name = 'Georgia'
        r_t.font.size = Pt(13)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(26, 58, 26)

        p_text = doc.add_paragraph()
        p_text.paragraph_format.space_after = Pt(10)
        r_txt = p_text.add_run(text)
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(10.5)
        r_txt.font.color.rgb = RGBColor(51, 51, 51)

    # Page Break for Signatures
    doc.add_page_break()

    # Signatures block title
    h_sig = doc.add_paragraph()
    h_sig.paragraph_format.space_before = Pt(14)
    h_sig.paragraph_format.space_after = Pt(12)
    r_sig = h_sig.add_run("IN WITNESS WHEREOF, the Parties have executed this Mutual Non-Disclosure Agreement as of the Effective Date.")
    r_sig.font.name = 'Georgia'
    r_sig.font.size = Pt(12)
    r_sig.font.bold = True
    r_sig.font.color.rgb = RGBColor(26, 58, 26)

    # Signatures table (1 row, 2 columns)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    
    # Left Column
    cell_left = sig_table.rows[0].cells[0]
    p_l = cell_left.paragraphs[0]
    p_l.paragraph_format.line_spacing = 1.3
    r_l = p_l.add_run(
        "For: SHAMBAIQ OPERATIONS LIMITED\n\n"
        "By: _______________________________\n\n"
        "Name: _____________________________\n\n"
        "Title: Director / Authorized Representative\n\n"
        "Date: _____________________________"
    )
    r_l.font.name = 'Arial'
    r_l.font.size = Pt(10.5)
    r_l.font.color.rgb = RGBColor(51, 51, 51)

    # Right Column
    cell_right = sig_table.rows[0].cells[1]
    p_r = cell_right.paragraphs[0]
    p_r.paragraph_format.line_spacing = 1.3
    r_r = p_r.add_run(
        "For: THE PARTNER ORGANIZATION\n\n"
        "Company Name: _____________________\n\n"
        "By: _______________________________\n\n"
        "Name: _____________________________\n\n"
        "Title: Authorized Signatory\n\n"
        "Date: _____________________________"
    )
    r_r.font.name = 'Arial'
    r_r.font.size = Pt(10.5)
    r_r.font.color.rgb = RGBColor(51, 51, 51)

    # Remove borders for signature table or style it
    # We can keep standard borders or let Word handle it.

    # Save DOCX
    docx_path = os.path.abspath("shambaiq_mutual_nda.docx")
    doc.save(docx_path)
    print(f"Mutual NDA DOCX saved to {docx_path}")
    return docx_path

def create_sacco_doc(logo_path):
    print("Building ShambaIQ SACCO Licensing Suite DOCX...")
    doc = Document()
    add_header_footer(doc)

    # 1. Logo
    doc.add_picture(logo_path, width=Inches(5.0))
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # 2. Title & Subtitle
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Agricultural SACCO Licensing Suite")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 58, 26) # Dark Green

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Smart Ag-Lending & Input Credit De-Risking for Kenyan SACCOs")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(200, 134, 10) # Gold

    # 3. Metadata box
    p_meta = doc.add_paragraph()
    set_paragraph_callout(p_meta, "F5F0E1", "8B6914") # Cream bg, brown border
    p_meta.paragraph_format.space_after = Pt(24)
    
    meta_text = (
        "Document Version: 1.0 (May 2026)\n"
        "Prepared by: ShambaIQ B2B Operations & Risk Advisory\n"
        "Official Email: info@shambaiq.com  |  Official WhatsApp: +254 748 042 633\n"
        "Principal Office: Nairobi, Kenya"
    )
    run_meta = p_meta.add_run(meta_text)
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(26, 58, 26)

    # 4. Heading 1: Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    run_h1 = h1.add_run("1. Executive Summary")
    run_h1.font.name = 'Georgia'
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = RGBColor(26, 58, 26)

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_after = Pt(12)
    run_exec = p_exec.add_run(
        "In Kenya, Savings and Credit Co-operative Societies (SACCOs) are the backbone of agricultural financing, "
        "providing critical credit for seeds, fertilizers, and machinery. However, agricultural loan books face "
        "substantial systemic risk: crop failure.\n\n"
        "A primary, yet preventable, cause of crop failure is poor agronomic decision-making. When a farmer applies the "
        "wrong fertilizer (such as standard DAP on highly acidic soils, which locks up phosphorus), crop yields plummet, "
        "and loan defaults skyrocket.\n\n"
        "ShambaIQ offers a data-driven solution. By integrating high-resolution satellite soil data (iSDAsoil) and "
        "localized crop suitability models, we enable SACCOs to de-risk agricultural lending at the point of disbursement. "
        "By verifying that input credit matches exact local soil chemistry, SACCOs protect their capital, increase farmer yields, "
        "and ensure high loan repayment rates."
    )
    run_exec.font.name = 'Arial'
    run_exec.font.size = Pt(11)
    run_exec.font.color.rgb = RGBColor(51, 51, 51)

    # 5. Heading 1: Core Offerings for SACCOs
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    run_h2 = h2.add_run("2. Core Offerings for SACCOs")
    run_h2.font.name = 'Georgia'
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = RGBColor(26, 58, 26)

    offerings = [
        ("Credit Underwriting API (B2B SaaS): ", "Your mobile banking app or USSD loan application channel queries our API with the farmer's location (GPS coordinates or county/ward) and target crop. Instantly returns a soil suitability grade and a recommended input budget, allowing you to set credit limits dynamically based on actual agronomic needs."),
        ("Co-Branded Credit Officer Portal: ", "A secure web portal hosted on a dedicated subdomain (e.g., sacco-advisory.shambaiq.com) used by credit officers during loan appraisals to generate a Soil Suitability & Fertilizer Prescription Report to attach to the loan file. Recommendations can be restricted strictly to pre-approved input partners."),
        ("Input Voucher Verification: ", "Integration with the SACCO's input check-off system or agrovet voucher program, ensuring that vouchers are automatically restricted to the exact fertilizer and seed varieties approved by ShambaIQ for that member's specific ward.")
    ]

    for title, desc in offerings:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(6)
        r_title = p_bullet.add_run(title)
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(26, 58, 26)
        r_desc = p_bullet.add_run(desc)
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(11)
        r_desc.font.color.rgb = RGBColor(51, 51, 51)

    # Page Break
    doc.add_page_break()

    # 6. Heading 1: Commercial Packages & Licensing Tiers
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(8)
    run_h3 = h3.add_run("3. Commercial Packages & Licensing Tiers")
    run_h3.font.name = 'Georgia'
    run_h3.font.size = Pt(16)
    run_h3.font.bold = True
    run_h3.font.color.rgb = RGBColor(26, 58, 26)

    doc.add_paragraph("We offer flexible pricing models structured around SACCO asset sizes and membership volumes:").runs[0].font.name = 'Arial'

    # Create Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Licensing Tier'
    hdr_cells[1].text = 'Best For & Focus'
    hdr_cells[2].text = 'Commercial Pricing (KES)'

    # Style Header row
    for cell in hdr_cells:
        set_cell_shading(cell, "16A34A")
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10.5)

    # Row 1
    row_cells = table.rows[1].cells
    row_cells[0].text = "Tier 1: Loan Appraisal API"
    row_cells[1].text = "Digital-first SACCOs with mobile loan platforms or USSD menus."
    row_cells[2].text = "• Basic Package: KES 35,000 / month (Up to 10k queries)\n• Enterprise Package: KES 85,000 / month (Up to 50k queries)\n• Volume Rate: KES 0.90 per query thereafter"

    # Row 2
    row_cells = table.rows[2].cells
    row_cells[0].text = "Tier 2: Co-Branded Portal License"
    row_cells[1].text = "Traditional SACCOs with physical branches and field extension officers."
    row_cells[2].text = "• Annual Base License: KES 180,000 / year (Includes branch-level access and customized input catalogs)\n• Setup & Integration Fee: KES 30,000 (one-time)"

    # Row 3
    row_cells = table.rows[3].cells
    row_cells[0].text = "Tier 3: Aggregated Risk & Yield Mapping"
    row_cells[1].text = "SACCO boards and risk committees seeking macro-portfolio insights."
    row_cells[2].text = "• Custom Regional Analysis: Starting at KES 100,000 per county\n• Delivery: Comprehensive spatial analysis mapping member farms against local soil chemistry metrics."

    # Style table text
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue
        bg_color = "F5F0E1" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, bg_color)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(51, 51, 51)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(18)

    # 7. Heading 1: Why ShambaIQ is the Best Partner
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(8)
    run_h4 = h4.add_run("4. Why ShambaIQ is the Best Partner")
    run_h4.font.name = 'Georgia'
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = RGBColor(26, 58, 26)

    reasons = [
        ("No Soil Testing Delays: ", "Traditional wet-chemistry soil tests take 7 to 14 days and cost KES 1,500+ per farm. ShambaIQ delivers instant sub-second digital diagnostics for pH, N, P, and K down to a 30m resolution."),
        ("Fraud Prevention: ", "Ensure that input credit disbursed is spent on agronomic assets matching the actual soil needs, reducing diversion of cash loans."),
        ("Zero Infrastructure Overhead: ", "We host and maintain the entire agronomic model. Your technical teams integrate via standard REST/JSON endpoints.")
    ]

    for title, desc in reasons:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(6)
        r_title = p_bullet.add_run(title)
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(11)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(26, 58, 26)
        r_desc = p_bullet.add_run(desc)
        r_desc.font.name = 'Arial'
        r_desc.font.size = Pt(11)
        r_desc.font.color.rgb = RGBColor(51, 51, 51)

    # 8. Heading 1: Get Started
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(8)
    run_h5 = h5.add_run("5. Get Started")
    run_h5.font.name = 'Georgia'
    run_h5.font.size = Pt(16)
    run_h5.font.bold = True
    run_h5.font.color.rgb = RGBColor(26, 58, 26)

    p_start = doc.add_paragraph()
    p_start.paragraph_format.space_after = Pt(12)
    run_start = p_start.add_run(
        "Protecting SACCO risk data and our agronomic algorithms is highly important. To request a demo portal login or schedule "
        "a presentation for your credit committee, please review our standard Mutual NDA and contact our partnerships team:"
    )
    run_start.font.name = 'Arial'
    run_start.font.size = Pt(11)
    run_start.font.color.rgb = RGBColor(51, 51, 51)

    # Callout for contact info
    p_contact = doc.add_paragraph()
    set_paragraph_callout(p_contact, "E8F5E9", "16A34A")
    p_contact.paragraph_format.space_after = Pt(12)
    
    contact_text = (
        "SHAMBAIQ SACCO RELATIONSHIP DIVISION\n"
        "• Primary Email: info@shambaiq.com\n"
        "• Risk Advisory Hotline: +254 748 042 633\n"
        "• Operations Center: Nairobi, Kenya"
    )
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = 'Arial'
    run_contact.font.size = Pt(11)
    run_contact.font.bold = True
    run_contact.font.color.rgb = RGBColor(26, 58, 26)

    # Save DOCX
    docx_path = os.path.abspath("shambaiq_sacco_proposal.docx")
    doc.save(docx_path)
    print(f"SACCO Suite DOCX saved to {docx_path}")
    return docx_path

def convert_to_pdf(word, docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    print(f"Converting {docx_path} to PDF...")
    doc = None
    try:
        doc = word.Documents.Open(docx_path)
        # 17 represents wdFormatPDF
        doc.SaveAs(pdf_path, FileFormat=17)
        print(f"Success: PDF generated at {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"Error during PDF conversion of {docx_path}: {e}")
        return None
    finally:
        if doc:
            try:
                doc.Close()
            except Exception:
                pass

if __name__ == "__main__":
    logo_path = generate_logo()
    
    prop_docx = create_partnership_doc(logo_path)
    nda_docx = create_nda_doc(logo_path)
    sacco_docx = create_sacco_doc(logo_path)
    
    print("\nStarting PDF conversion via MS Word...")
    word = None
    try:
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        
        convert_to_pdf(word, prop_docx)
        convert_to_pdf(word, nda_docx)
        convert_to_pdf(word, sacco_docx)
    except Exception as e:
        print(f"Failed to initialize MS Word COM server: {e}")
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
    
    print("\nAll documents generated successfully!")
